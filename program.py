
from tkinter import *
from tkinter import filedialog as fd

import matplotlib
import numpy as np
import pandas as pd

matplotlib.use("TkAgg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

import math
from sklearn.preprocessing import MinMaxScaler

# import tensorflow
# import keras
from keras.models import Sequential
from keras.layers.core import Dense
from keras.layers.recurrent import LSTM

import calendar
from datetime import timedelta, datetime


# import lxml
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # выравнивание


def insert_csv():                               # Выбираем и загружаем нужный файл в pandas таблицу
    global dataframe, date_doc
    dataframe = pd.read_csv(fd.askopenfilename(), header=0, sep=";")        # , usecols=[1]
    date_doc = dataframe.iloc[-1, 0]
    dataframe = dataframe.iloc[:,1]

    Button_check.config(state='normal')


def insert_doc():
    global dataframe, report_n, report_date, date_doc
    wordDoc = Document(fd.askopenfilename())

    table_size = len(wordDoc.tables[1].rows)
    report_n = wordDoc.tables[0].rows[-1].cells[0].text
    report_date = wordDoc.tables[0].rows[-1].cells[1].text
    date_doc = wordDoc.tables[1].rows[table_size - 1].cells[0].text

    dataframe = []
    for i in range(1, table_size):
        dataframe.append(wordDoc.tables[1].rows[i].cells[1].text)
    dataframe = pd.DataFrame(dataframe)

    Button_check.config(state='normal')


def extract_csv():
    b = pd.DataFrame(df_prognoz)
    b.to_csv(fd.asksaveasfilename(defaultextension=".csv"), sep=";", header=0)


def extract_doc():
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7874)
        section.bottom_margin = Inches(0.7874)
        section.left_margin = Inches(1.1811)
        section.right_margin = Inches(0.5905)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    styles = document.styles
    charstyle = styles.add_style('Small', style_type=1)  # CHARACTER
    font = charstyle.font
    font.size = Pt(10)
    font.name = 'Times New Roman'


    document.add_paragraph('УТВЕРЖДАЮ').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('Начальник отдела планирования').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('_____________________').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('«__» ____________201_г').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()
    document.add_paragraph('___________________________________________________________'
                           ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('(наименование предприятия)', style='Small').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('ПРОГНОЗ ОБЪЕМА ПРОДАЖ').alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер документа'
    hdr_cells[1].text = 'Дата составления прогноза'
    hdr_cells[2].text = 'На основе отчета объема продаж для прогнозирования'
    hdr_cells[3].text = 'Дата составления отчета'
    hdr_cells[4].text = 'Средний процент ошибок на тесте'

    row_cells = table.add_row().cells
    row_cells[0].text = ''
    now = datetime.now()
    row_cells[1].text = now.strftime("%d-%m-%Y")
    row_cells[2].text = report_n
    row_cells[3].text = report_date
    row_cells[4].text = testmape

    document.add_paragraph()

    table1 = document.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text = 'Месяц, год'
    hdr_cells1[1].text = 'Значение'
    for i in range(prognoz_size):
        row_cells1 = table1.add_row().cells
        row_cells1[0].text = df_prognoz[i, 0]
        row_cells1[1].text = str(df_prognoz[i, 1])          # round()   a = "%.2f" % a

    document.add_paragraph()
    document.add_paragraph('Сотрудник отдела планирования	__________		_____________________'
                           ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(' (подпись)		(расшифровка подписи)		',
                           style='Small').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save(fd.asksaveasfilename(defaultextension=".docx"))


def TRAIN_XY(look_back, train):
    dataX, dataY = [], []                                           # разделяем на Х и Y для теста
    for i in range(len(train)-look_back):
        dataX.append(train[i:i+look_back])
        dataY.append(train[i+look_back,0])

    trainX = np.array(dataX)
    trainY = np.array(dataY)
    trainX = trainX.reshape(trainX.shape[0], 1, trainX.shape[1])    # изменяем размер для иксов

    return (trainX, trainY)


def TRAIN_MODEL (trainX, trainY, look_back, train_len):

    model = Sequential()                                                        # Creating a model
    model.add(LSTM(6, input_shape=(1, look_back), activation='sigmoid'))
    model.add(Dense(1, activation='linear'))

    model.compile(loss='mean_squared_error',                                    # Compiling model
                  optimizer='adam',
                  metrics=['mean_absolute_percentage_error'])

    model.fit(trainX, trainY, batch_size=1, epochs=train_len, verbose=2)           # Training a model

    return model


def PREDICT(X, size):

    Predict = np.array([None], ndmin=2)

    for i in range(size):
        elpredictions = model.predict(X)                              # вычисляю новое предсказание
        Predict = np.concatenate([Predict, elpredictions])            # записываю предсказание
        X = X[:, :, 1:]                                               # отбрасываю 1й элемент
        X = np.concatenate((X, elpredictions), axis=None)             # беру в строку предсказанный элемент
        X = np.array(X, ndmin=2)
        X = X.reshape(X.shape[0], 1, X.shape[1])

    Predict = Predict[1:]  # Отбрасываю первую пустую строку

    return Predict


def TEST_MAPE(testY, testPredict, test_size):
    testmape = 0
    for i in range(test_size):
        mapetestEl = math.fabs((int(testY[i, :]) - int(testPredict[i, :])) / testY[i, :])
        testmape += mapetestEl
    testmape = (testmape / test_size) * 100
    testmape = "%.2f" % testmape
    return (testmape)


def PROGNOZ_DATE(date_doc, prognoz_size, prognozPredict):

    b = date_doc
    b = datetime.strptime(date_doc, '%Y-%m')
    df_prognoz = np.array([])
    for i in range(1, prognoz_size + 1):
        days_in_month = calendar.monthrange(b.year, b.month)[1]
        b += timedelta(days=days_in_month)
        df_prognoz = np.concatenate([df_prognoz, [b.strftime('%Y-%m')]])
    df_prognoz = df_prognoz.reshape(-1, 1)
    df_prognoz = np.concatenate([df_prognoz, prognozPredict], axis=1)

    return df_prognoz


def GRAPH(dataset, scaler, trainPredictPlot = [None], testPredictPlot = [None], prognozPredictPlot = [None]):

    Figure_graph = Figure(figsize=(8, 5), dpi=100)

    Plot = Figure_graph.add_subplot(111)
    Plot.plot(scaler.inverse_transform(dataset), 'b')
    Plot.plot(trainPredictPlot, 'g')
    Plot.plot(testPredictPlot, 'r')
    Plot.plot(prognozPredictPlot, 'c')
    Plot.grid()

    Canvas = FigureCanvasTkAgg(Figure_graph, master=window)
    Canvas.draw()
    Canvas.get_tk_widget().grid(row=0, column=1, rowspan=30)


def CHECK():
    global dataset, scaler
    # Canvas_space.get_tk_widget().destroy()

    dataset_not_normalize = dataframe.values  # переводим из pandas в numpy
    dataset_not_normalize = dataset_not_normalize.astype('float32')  # приводим значения к типу float
    dataset_not_normalize = dataset_not_normalize.reshape(-1, 1)

    scaler = MinMaxScaler(feature_range=(0, 1))  # normalize the dataset
    dataset = scaler.fit_transform(dataset_not_normalize)

    GRAPH(dataset, scaler)

    Button_train.config(state='normal')


def TRAIN():
    global look_back, prognoz_size, test_size, \
        model, train, test, trainPredictPlot

    look_back = int(Entry_look_back.get())              # получаем значения из entry
    prognoz_size = int(Entry_prognoz_size.get())
    test_size = int(Entry_test_size.get())
    train_len = int(Entry_train_len.get())

    train_size = len(dataset) - test_size               # вычисляем обучающую выборку
    train = dataset[:train_size, ]                       # делим на тестовую и обучающую
    test = dataset[train_size:len(dataset), :]

    # получаем значения для работы сети
    trainX, trainY = TRAIN_XY(look_back, train)


    # тренируем модель
    model = TRAIN_MODEL(trainX, trainY, look_back, train_len)


    # make predictions
    trainPredict = model.predict(trainX)

    # invert predictions
    trainPredict = scaler.inverse_transform(trainPredict)

    # shift train predictions for plotting
    trainPredictPlot = np.empty_like(dataset)
    trainPredictPlot[:, :] = np.nan
    trainPredictPlot[look_back:len(trainPredict) + look_back, :] = trainPredict


    # строим график
    GRAPH(dataset, scaler, trainPredictPlot)

    Button_test.config(state='normal')


def TEST():
    global testPredictPlot, testPredict, testmape

    testX = dataset[-(look_back+test_size):-test_size]
    testX = testX.reshape(1, testX.shape[1], testX.shape[0])
    testY = test

    testPredict = PREDICT(testX, test_size)                                     # получаем предсказания

    testPredict = scaler.inverse_transform(testPredict)                         # invert predictions

    testPredictPlot = np.empty_like(dataset)                                    # shift test predictions for plotting
    testPredictPlot[:, :] = np.nan
    testPredictPlot[-test_size:len(dataset), :] = testPredict

    GRAPH(dataset, scaler, trainPredictPlot, testPredictPlot)                   # строим график

    testY = scaler.inverse_transform(testY)
    testmape = TEST_MAPE(testY, testPredict, test_size)

    Ltestmape.configure(text="MAPE = " + testmape)

    Button_prognoz.config(state='normal')


def PROGNOZ():
    global prognozPredict, df_prognoz

    Entry_look_back.config(state='disabled')
    Entry_prognoz_size.config(state='disabled')
    Entry_test_size.config(state='disabled')
    Entry_train_len.config(state='disabled')

    prognozX = dataset[-look_back:]
    prognozX = prognozX.reshape(1, prognozX.shape[1], prognozX.shape[0])

    prognozPredict = PREDICT(prognozX, prognoz_size)                    # получение предсказаний

    prognozPredict = scaler.inverse_transform(prognozPredict)           # invert predictions

    datasetprognoz = prognoz_size * [None]                              # create datasetprognoz
    datasetprognoz = np.array(datasetprognoz, ndmin=2)
    datasetprognoz = np.transpose(datasetprognoz)
    datasetprognoz = np.concatenate((dataset, datasetprognoz))

    prognozPredictPlot = np.empty_like(datasetprognoz)                  # shift test predictions for plotting
    prognozPredictPlot[:, :] = np.nan
    prognozPredictPlot[len(dataset):len(dataset) + prognoz_size, :] = prognozPredict

    df_prognoz = PROGNOZ_DATE(date_doc, prognoz_size, prognozPredict)

    GRAPH(dataset, scaler, trainPredictPlot, testPredictPlot, prognozPredictPlot)

    Button_extract_csv.config(state='normal')
    Button_extract_doc.config(state='normal')


# Создаем окно приложения
window = Tk()
window.title("Прогнозирование объема продаж")

# Создание нужных переменных
dataframe = pd.DataFrame()
look_back = int()
prognoz_size = int()
test_size = int()


# Создания фрейма под меню
Frame_menu = Frame(window)
Frame_menu.grid(row=0, column=0, pady=5, padx=10)


# Создаем и размещаем виджет для записи размера окна
Llook_back = Label(Frame_menu, text="Размер окна")
Llook_back.grid(row=2, column=0, sticky=W, pady=2)
Entry_look_back = Entry(Frame_menu)   # textvariable = look_back, variable=look_back
Entry_look_back.grid(row=2, column=1, sticky=E, pady=2)

# Создаем и размещаем виджет для записи размера прогноза
Lprognoz_size = Label(Frame_menu, text="Размер прогноза")
Lprognoz_size.grid(row=3, column=0, sticky=W, pady=2)
Entry_prognoz_size = Entry(Frame_menu)
Entry_prognoz_size.grid(row=3, column=1, sticky=E, pady=2)

# Создаем и размещаем виджет для записи размера тестовой выборки
Ltest_size = Label(Frame_menu, text="Размер тестовой выборки")
Ltest_size.grid(row=4, column=0, sticky=W, pady=2)
Entry_test_size = Entry(Frame_menu)
Entry_test_size.grid(row=4, column=1, pady=2)

# Создаем и размещаем виджет для записи длительности обучения
Ltrain_len = Label(Frame_menu, text="Длительность обучения")
Ltrain_len.grid(row=5, column=0, sticky=W, pady=2)
Entry_train_len = Entry(Frame_menu)
Entry_train_len.grid(row=5, column=1, pady=2)


# Создаем кнопку для выбора файла csv
Button_path_csv = Button(Frame_menu, text="Открыть csv", command=insert_csv, width=11)
Button_path_csv.grid(row=0, column=0, pady=2)

# Создаем кнопку для выбора файла docx
Button_path_doc = Button(Frame_menu, text="Открыть doc", command=insert_doc, width=11)
Button_path_doc.grid(row=0, column=1, pady=2)

# Создаем кнопку для проверки данных
Button_check = Button(Frame_menu, text="Проверка", command=CHECK, state='disabled', width=11)
Button_check.grid(row=1, column=1, pady=2)

# Создаем кнопку для начала тренировки сети
Button_train = Button(Frame_menu, text="Обучение", command=TRAIN, state='disabled', width=11)
Button_train.grid(row=6, column=1, pady=2)

# Создаем кнопку для начала теста
Button_test = Button(Frame_menu, text="Тестирование", command=TEST, state='disabled', width=11)
Button_test.grid(row=7, column=1, pady=2)

# Создаем кнопку для вычисления прогноза
Button_prognoz = Button(Frame_menu, text="Прогноз", command=PROGNOZ, state='disabled', width=11)
Button_prognoz.grid(row=8, column=1, pady=2)

# Создаем кнопку для вывода результатов в csv
Button_extract_csv = Button (Frame_menu, text="Сохранить csv", command=extract_csv, state='disabled', width=11)
Button_extract_csv.grid(row=9, column=0, pady=2)

# Создаем кнопку для вывода результатов в docx
Button_extract_doc = Button (Frame_menu, text="Сохранить doc", command=extract_doc, state='disabled', width=11)
Button_extract_doc.grid(row=9, column=1, pady=2)

# Создаем виджет для вывода ошибки
Ltestmape = Label(Frame_menu, text= "MAPE =" )
Ltestmape.grid(row=7, column=0, sticky=W, pady=2)


Frame_graph = Frame(window)
Frame_graph.grid(row=0, column=1)

Figure_space = Figure(figsize=(8, 5), dpi=100)

Canvas_space = FigureCanvasTkAgg(Figure_space, master=window)
Canvas_space.get_tk_widget().grid(row=0, column=1, rowspan=30)



window.mainloop()
